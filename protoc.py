from docx import Document
from os import system

DOCUMENT_NAME = "Protokoll.docx"

document = Document('Protokoll.docx')

eins = document.add_paragraph("Überschrift 1")
zwei = document.add_paragraph("Überschrift 2")
drei = document.add_paragraph("Überschrift 3")


"""
zwei.insert_paragraph_before("do gheats eini", style = 'List Bullet')
zwei.insert_paragraph_before("zweiter Teil", style = 'List Bullet')
drei.insert_paragraph_before("text für zwei", style = 'List Bullet')
"""

dicti = {"zwei" : zwei}
while True:
    eingabe = input("paragraph: ")
    if eingabe != "aus":
        txt = input("text: ")
        if txt != "aus":
            dicti[eingabe].insert_paragraph_before(txt, style = 'List Bullet')
            #zwei.add_run("fixbestand")
            document.save(DOCUMENT_NAME)
    else:
        break
    
    
    
    

document.save(DOCUMENT_NAME)

system(f'open {DOCUMENT_NAME}')
